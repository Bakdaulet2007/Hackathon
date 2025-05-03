import os
from flask import (
    Flask, request, render_template, redirect,
    url_for, send_from_directory, flash
)
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, UserMixin,
    login_user, login_required,
    logout_user, current_user
)
from werkzeug.security import generate_password_hash, check_password_hash
import google.generativeai as genai
import pdfplumber
from docx import Document

# ——————————————————————————————————————————————
# 1) Настройка Gemini API
genai.configure(api_key="AIzaSyAhlYLJRhT7Cyh_ViBq5ppT89xN_9-fNgo")

# 2) Инициализация Flask и расширений
app = Flask(__name__)
app.config['SECRET_KEY'] = 'замени_на_сложный_секрет'

# постарайтесь хранить БД рядом с этим файлом
basedir = os.path.abspath(os.path.dirname(__file__))
db_path = os.path.join(basedir, 'site.db')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + db_path
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

print(f">>> БД будет создана здесь: {db_path}")

# ——————————————————————————————————————————————
# 3) Модели
class User(db.Model, UserMixin):
    id       = db.Column(db.Integer, primary_key=True)
    email    = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(128), nullable=False)
    role     = db.Column(db.String(20), default='student')  # teacher или student

class Lecture(db.Model):
    id               = db.Column(db.Integer, primary_key=True)
    filename         = db.Column(db.String, nullable=False)
    summary_text     = db.Column(db.Text,   nullable=False)
    summary_filename = db.Column(db.String, nullable=False)
    uploaded_by      = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    timestamp        = db.Column(db.DateTime, server_default=db.func.now())

# 4) Создаем таблицы сразу после объявления моделей
with app.app_context():
    print(">>> Создаём таблицы…")
    db.create_all()
    print(">>> Таблицы созданы.")

# 5) Flask-Login: загрузка пользователя
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# ——————————————————————————————————————————————
# 6) Утилиты для текстов и конспектов
def extract_text(path, ext):
    if ext == '.pdf':
        text = ''
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ''
        return text
    elif ext == '.docx':
        doc = Document(path)
        return '\n'.join(p.text for p in doc.paragraphs)
    return ''

def get_summary(text):
    model = genai.GenerativeModel(model_name="models/gemini-1.5-pro")
    prompt = f"""
Напиши структурированный конспект по следующей лекции:
1) Разбей на разделы с понятными заголовками.
2) В каждом разделе выдели ключевые идеи списком.
3) Добавь определения важных терминов.
4) Приведи примеры.
Текст лекции:
{text}
"""
    response = model.generate_content([prompt])
    return response.text

def save_summary_docx(summary, original_name):
    doc = Document()
    doc.add_heading('Конспект', level=1)
    for line in summary.split('\n'):
        doc.add_paragraph(line)
    base  = os.path.splitext(original_name)[0]
    fname = f"{base}_summary.docx"
    os.makedirs('uploads', exist_ok=True)
    out_path = os.path.join('uploads', fname)
    doc.save(out_path)
    return fname

# ——————————————————————————————————————————————
# 7) Регистрация и логин
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email    = request.form['email']
        pwd_hash = generate_password_hash(request.form['password'])
        role     = request.form.get('role', 'student')
        user     = User(email=email, password=pwd_hash, role=role)
        db.session.add(user)
        db.session.commit()
        login_user(user)
        return redirect(url_for('index'))
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        user = User.query.filter_by(email=request.form['email']).first()
        if user and check_password_hash(user.password, request.form['password']):
            login_user(user)
            # Перенаправляем по роли
            if user.role == 'teacher':
                return redirect(url_for('index'))
            else:
                return redirect(url_for('student'))
        flash('Неверный логин или пароль', 'danger')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

# ——————————————————————————————————————————————
# 8) Интерфейс преподавателя
@app.route('/', methods=['GET'])
@login_required
def index():
    # Если это студент — сразу на его страницу
    if current_user.role == 'student':
        return redirect(url_for('student'))
    # Иначе — преподаватель
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if current_user.role != 'teacher':
        return "Доступ только для преподавателей", 403

    file = request.files.get('file')
    if not file or not file.filename:
        return "Файл не выбран", 400

    orig_name = file.filename
    ext       = os.path.splitext(orig_name)[1].lower()
    if ext not in ('.pdf', '.docx'):
        return "Только PDF или DOCX поддерживаются", 400

    os.makedirs('uploads', exist_ok=True)
    path = os.path.join('uploads', orig_name)
    file.save(path)

    text          = extract_text(path, ext)
    summary       = get_summary(text)
    summary_fname = save_summary_docx(summary, orig_name)

    lecture = Lecture(
        filename         = orig_name,
        summary_text     = summary,
        summary_filename = summary_fname,
        uploaded_by      = current_user.id
    )
    db.session.add(lecture)
    db.session.commit()

    download_url = url_for('download_file', filename=summary_fname)
    return f"""
      <h2>Конспект готов!</h2>
      <p><a href="{download_url}">Скачать конспект (.docx)</a></p>
    """

# ——————————————————————————————————————————————
# 9) Интерфейс студента
@app.route('/student', methods=['GET'])
@login_required
def student():
    if current_user.role != 'student':
        return "Доступ только для студентов", 403
    lectures = Lecture.query.order_by(Lecture.timestamp.desc()).all()
    return render_template('student.html', lectures=lectures)

# ——————————————————————————————————————————————
# 10) Раздача файлов
@app.route('/uploads/<path:filename>')
def download_file(filename):
    return send_from_directory('uploads', filename, as_attachment=True)

# ——————————————————————————————————————————————
if __name__ == '__main__':
    app.run(debug=True)
